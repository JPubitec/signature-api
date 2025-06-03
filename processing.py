import io
import os
import base64
from cryptography.hazmat.primitives import hashes, serialization
from cryptography.hazmat.primitives.asymmetric import ec
from cryptography.hazmat.primitives.asymmetric.utils import Prehashed
from cryptography.hazmat.backends import default_backend
from cryptography.exceptions import InvalidSignature
from docx import Document

# =======================
# GENERACIÓN DE CLAVES
# =======================

def generar_claves(path_priv="clave_privada.pem", path_pub="clave_publica.pem"):
    if os.path.exists(path_priv) and os.path.exists(path_pub):
        print("🛑 Las claves ya existen. No se generaron nuevas.")
        return

    print("🔐 Generando nuevas claves privada y pública...")
    private_key = ec.generate_private_key(ec.SECP256R1())
    public_key = private_key.public_key()

    with open(path_priv, "wb") as f:
        f.write(private_key.private_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PrivateFormat.PKCS8,
            encryption_algorithm=serialization.NoEncryption()
        ))

    with open(path_pub, "wb") as f:
        f.write(public_key.public_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PublicFormat.SubjectPublicKeyInfo
        ))

    print("✅ Claves generadas y guardadas.")

# =======================
# CARGA DE CLAVES
# =======================

def cargar_clave_privada(path="clave_privada.pem"):
    with open(path, "rb") as f:
        return serialization.load_pem_private_key(
            f.read(),
            password=None,
            backend=default_backend()
        )

def cargar_clave_publica(path="clave_publica.pem"):
    with open(path, "rb") as f:
        return serialization.load_pem_public_key(
            f.read(),
            backend=default_backend()
        )

# =======================
# FIRMA DE ARCHIVO (BYTES)
# =======================

def firmar_docx(docx_bytes, private_key):
    digest = hashes.Hash(hashes.SHA256())
    digest.update(docx_bytes)
    hash_final = digest.finalize()

    firma = private_key.sign(
        hash_final,
        ec.ECDSA(Prehashed(hashes.SHA256()))
    )

    print("✔ Documento firmado correctamente.")
    return firma

# =======================
# VERIFICACIÓN DE FIRMA (BYTES + BASE64)
# =======================

def verificar_firma_docx(docx_bytes, firma_b64, public_key):
    firma = base64.b64decode(firma_b64)

    digest = hashes.Hash(hashes.SHA256())
    digest.update(docx_bytes)
    hash_final = digest.finalize()

    try:
        public_key.verify(
            firma,
            hash_final,
            ec.ECDSA(Prehashed(hashes.SHA256()))
        )
        print("✔ Firma válida: el archivo es auténtico.")
        return True
    except InvalidSignature:
        print("✖ Firma inválida: el archivo ha sido modificado o la clave es incorrecta.")
        return False

# =======================
# INSERCIÓN DE FIRMA EN DOCX
# =======================

def insertar_firma_en_docx(docx_bytes, firma_b64):
    doc = Document(io.BytesIO(docx_bytes))
    doc.add_paragraph("FIRMA DIGITAL (Base64):")
    doc.add_paragraph(firma_b64)

    salida = io.BytesIO()
    doc.save(salida)
    return salida.getvalue()

# =======================
# EJEMPLO DE USO LOCAL
# =======================

if __name__ == "__main__":
    generar_claves()

    private_key = cargar_clave_privada()
    public_key = cargar_clave_publica()

    with open("Test.docx", "rb") as f:
        contenido_docx = f.read()

    firma = firmar_docx(contenido_docx, private_key)
    firma_b64 = base64.b64encode(firma).decode()

    verificar_firma_docx(contenido_docx, firma_b64, public_key)

    docx_firmado = insertar_firma_en_docx(contenido_docx, firma_b64)
    with open("Test_firmado.docx", "wb") as f:
        f.write(docx_firmado)
